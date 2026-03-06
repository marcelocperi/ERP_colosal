from docx import Document
from docx.shared import Pt
import os

def create_tooling_doc():
    doc = Document()
    doc.add_heading('ERP Colosal - Documentación de Utilitarios de Integridad', 0)
    doc.add_heading('Introducción', level=1)
    doc.add_paragraph('Herramientas de automatización para integridad de código y seguimiento de incidentes.')
    
    doc.add_heading('1. python_check_remediate.py', level=1)
    doc.add_paragraph('Audita y reconcilia discrepancias entre Menú, Blueprints y Templates.')
    doc.add_paragraph('Funciones:', style='List Bullet')
    doc.add_paragraph('Detección de rutas huérfanas.', style='List Bullet')
    doc.add_paragraph('Auto-generación de templates y controladores faltantes.', style='List Bullet')
    doc.add_paragraph('Registro de anomalías en sys_transaction_logs.', style='List Bullet')
    
    doc.add_heading('2. python_check_incident.py', level=1)
    doc.add_paragraph('Interfaz de consola para gestionar incidentes de la base de datos.')
    doc.add_paragraph('Funciones:', style='List Bullet')
    doc.add_paragraph('Listado de últimos incidentes.', style='List Bullet')
    doc.add_paragraph('Visualización de detalles técnicos por ID.', style='List Bullet')
    
    doc.save('erp_colosal_tooling.docx')

def create_masterplan_doc():
    doc = Document()
    doc.add_heading('Plan Maestro de Costos Industriales, Sourcing & R&D (MSAC v4.1)', 0)
    
    # Read the content from the MD file and try to format it simply
    if os.path.exists('master_plan_msac.md'):
        with open('master_plan_msac.md', 'r', encoding='utf-8') as f:
            lines = f.readlines()
            for line in lines:
                line = line.strip()
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('|'):
                    # Simple table representation or just paragraph
                    doc.add_paragraph(line)
                elif line:
                    doc.add_paragraph(line)
    
    doc.save('master_plan_msac.docx')

if __name__ == "__main__":
    try:
        create_tooling_doc()
        print("✅ erp_colosal_tooling.docx generado.")
        create_masterplan_doc()
        print("✅ master_plan_msac.docx generado.")
    except Exception as e:
        print(f"❌ Error: {e}")
