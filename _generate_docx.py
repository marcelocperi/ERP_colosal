from docx import Document
from docx.shared import Pt
import os

def create_doc():
    doc = Document()
    doc.add_heading('ERP Colosal - Documentación de Utilitarios de Integridad', 0)

    # Introducción
    doc.add_heading('Introducción', level=1)
    doc.add_paragraph(
        'Este documento describe las herramientas de automatización implementadas para asegurar la integridad '
        'del código, la consistencia de las rutas y el seguimiento de incidentes técnicos en el ERP Colosal.'
    )

    # python_check_remediate.py
    doc.add_heading('1. python_check_remediate.py', level=1)
    doc.add_paragraph(
        'Es el motor principal de reconciliación. Su función es auditar el ecosistema de la aplicación '
        'buscando discrepancias entre tres fuentes principales:'
    )
    doc.add_paragraph('Menu Structure (.agent/menu_structure.json): Rutas definidas para el usuario.', style='List Bullet')
    doc.add_paragraph('Código Python (routes.py): Implementación real de los controladores.', style='List Bullet')
    doc.add_paragraph('Templates HTML: Archivos físicos de vista.', style='List Bullet')
    
    doc.add_heading('Funciones Principales:', level=2)
    doc.add_paragraph('Detección de Rutas Huérfanas: Identifica rutas en Python que no están en el menú.', style='List Bullet')
    doc.add_paragraph('Validación de Templates: Verifica que cada ruta tenga su archivo HTML correspondiente.', style='List Bullet')
    doc.add_paragraph('Auto-Remediación: Crea automáticamente templates placeholder y controladores Python faltantes.', style='List Bullet')
    doc.add_paragraph('Consolidación de Menú: Corrige automáticamente nombres de Blueprints y elimina duplicados en el JSON del menú.', style='List Bullet')
    
    # python_check_incident.py
    doc.add_heading('2. python_check_incident.py', level=1)
    doc.add_paragraph(
        'Este utilitario permite la gestión y visualización de los incidentes generados por el proceso de reconciliación '
        'directamente desde la consola de comandos (PowerShell/CMD).'
    )
    
    doc.add_heading('Funciones Principales:', level=2)
    doc.add_paragraph('Listado de Incidentes: Muestra un resumen de los últimos incidentes registrados en sys_transaction_logs.', style='List Bullet')
    doc.add_paragraph('Detalle de Incidente: Permite ver el reporte técnico completo de un incidente específico mediante su ID.', style='List Bullet')
    doc.add_paragraph('Cumplimiento de Reglas: Asegura que no se generen archivos de log físicos, centralizando todo en la base de datos.', style='List Bullet')

    # Guardar
    path = os.path.join(os.getcwd(), 'erp_colosal_tooling.docx')
    doc.save(path)
    return path

if __name__ == "__main__":
    try:
        final_path = create_doc()
        print(f"✅ Documento generado exitosamente en: {final_path}")
    except Exception as e:
        print(f"❌ Error al generar el documento: {e}")
