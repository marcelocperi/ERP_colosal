
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"File {md_path} not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Estilo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    table_data = []
    in_table = False

    for line in lines:
        line = line.strip()
        
        # Headers
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            continue
        
        # Tables (Simplified MD table parser)
        if '|' in line:
            # Check if it's a separator line | --- | --- |
            if re.match(r'^[\s|:-]+$', line):
                continue
            
            # Row data
            row = [cell.strip() for cell in line.split('|') if cell.strip() or (cell == '' and '|' in line)]
            # If it's a real row (actually sometimes split leaves empty strings at start/end)
            # MD tables usually have | at ends: | cell1 | cell2 |
            parts = [p.strip() for p in line.split('|')]
            if len(parts) > 1:
                # Remove empty first/last if table starts/ends with |
                if parts[0] == '': parts.pop(0)
                if parts[-1] == '': parts.pop()
                
                if not in_table:
                    in_table = True
                    table_data = [parts]
                else:
                    table_data.append(parts)
            continue
        else:
            if in_table:
                # Flush table
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(table_data):
                        for j, cell_text in enumerate(row_data):
                            if j < len(table.columns):
                                # Clean bold/italic markers
                                clean_text = cell_text.replace('**', '').replace('✅', '✓').replace('⏳', '...').replace('❌', 'X')
                                table.cell(i, j).text = clean_text
                in_table = False
                table_data = []
            
            if not line:
                doc.add_paragraph('')
                continue
            
            # Simple bold/italic removal for paragraphs
            p_text = line.replace('**', '').replace('###', '').replace('---', '')
            if p_text.startswith('>'):
                p = doc.add_paragraph(p_text[1:].strip())
                p.style = 'Quote' if 'Quote' in doc.styles else 'Normal'
            else:
                doc.add_paragraph(p_text)

    # Final table flush if EOF
    if in_table and table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                if j < len(table.columns):
                    table.cell(i, j).text = cell_text.replace('**', '')

    doc.save(docx_path)
    print(f"Successfully converted to {docx_path}")

if __name__ == "__main__":
    md_file = "C:/Users/marce/Documents/GitHub/bibliotecaweb/multiMCP/master_plan_msac.md"
    docx_file = "C:/Users/marce/Documents/GitHub/bibliotecaweb/multiMCP/master_plan_msac.docx"
    md_to_docx(md_file, docx_file)
