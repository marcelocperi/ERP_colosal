import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_plan_docx():
    doc = Document()

    # Estilos del documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Título principal
    title = doc.add_heading('Plan Maestro de Avance y Costos (MSAC v4.2)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Generado automáticamente por el Sistema de Seguimiento Colosal ERP.\n\nEste documento presenta el estado de avance actual de la arquitectura industrial, los hitos consolidados, y la proyección de esfuerzo/costo para las fases subsiguientes (estimado en horas de desarrollo y consultoría experta).')

    # FASE 1
    h1 = doc.add_heading('Fase 1: Estructura Industrial, BOM & Estándares', level=1)
    doc.add_paragraph('Estado: COMPLETADA (100%)\nCosto Pendiente: 0 hrs (Fase Consolidada y Entregada)', style='Intense Quote')
    
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Light Shading Accent 1'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Sub-Fase / Módulo'
    hdr_cells[1].text = 'Estado'
    hdr_cells[2].text = 'Impacto / Resumen'

    items_fase_1 = [
        ('1.1 Maestro Multi-Origen', 'Completado', 'Compras Locales, Importadas, Producción y Fazón.'),
        ('1.2 Bill of Materials (BOM)', 'Completado', 'Recetas recursivas y explosión multi-nivel.'),
        ('1.3 Carga de Costos Indirectos', 'Completado', 'Overhead y Mano de Obra vía Standard Costing.'),
        ('1.4 Módulo Consignación & Fazón', 'Completado', 'Estructura de base de datos para Talleres Externos.'),
        ('1.5 Calidad & Técnico Legal', 'Completado', 'Repositorio Polimórfico (RNE, RNPA, ISO).'),
        ('1.6 Auditoría e Integridad SQL', 'Completado', 'Reconciliación de permisos y estructura SQL masiva.'),
        ('1.7 Tunning Experto BD', 'Completado', 'Generación de +400 índices para alto volumen.')
    ]

    for item, estado, impacto in items_fase_1:
        row_cells = table1.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = estado
        row_cells[2].text = impacto

    doc.add_page_break()

    # FASE 2
    doc.add_heading('Fase 2: Proyectos R&D y Trazabilidad Transformacional', level=1)
    doc.add_paragraph('Estado: EN EJECUCIÓN\nCosto Pendiente Restante: 120 hrs', style='Intense Quote')

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Light Shading Accent 1'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Sub-Fase / Módulo'
    hdr_cells2[1].text = 'Estado'
    hdr_cells2[2].text = 'Costo Estimado'
    hdr_cells2[3].text = 'Detalle Técnico'

    items_fase_2 = [
        ('2.1 Módulo Proyectos de Desarrollo', 'Completado', '0 hrs', 'Maestro I+D, Gastos de Ingeniería y Fases (Evaluación a Aprobado) ya en Base de Datos.'),
        ('2.2 Interfaz Fazón & Liquidación', 'Pendiente', '50 hrs', 'Desarrollo de pantallas (Front-end/Routes) para despacho a Taller Externo y liquidación de consumo vs devoluciones.'),
        ('2.3 Layer Roll-up (Producción)', 'Pendiente', '70 hrs', 'Desarrollo del motor de cierre productivo. Capitaliza costo de materiales FIFO + Overhead + Calidad.')
    ]

    for item, estado, costo, detalle in items_fase_2:
        row_cells = table2.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = estado
        row_cells[2].text = costo
        row_cells[3].text = detalle

    doc.add_paragraph('\n')

    # FASE 3
    doc.add_heading('Fase 3: Valuation, Margins & Finanzas', level=1)
    doc.add_paragraph('Estado: PENDIENTE\nCosto Pendiente Restante: 160 hrs', style='Intense Quote')

    table3 = doc.add_table(rows=1, cols=4)
    table3.style = 'Light Shading Accent 1'
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = 'Sub-Fase / Módulo'
    hdr_cells3[1].text = 'Estado'
    hdr_cells3[2].text = 'Costo Estimado'
    hdr_cells3[3].text = 'Detalle Técnico'

    items_fase_3 = [
        ('3.1 FIFO/WAC Across Layers', 'Pendiente', '60 hrs', 'Módulo financiero de recálculo de costos en cascada para productos terminados basados en el costo real del insumo importado o fabricado.'),
        ('3.2 Dynamic Margin Engine', 'Pendiente', '60 hrs', 'Pantalla de fijación y simulación de precios. Costo + Margen Bruto vs Neto + Impuestos.'),
        ('3.3 Pricing Synchronization', 'Pendiente', '40 hrs', 'Jobs Asíncronos. Alertas cuando un proveedor varía el precio impactando en la rentabilidad industrial.')
    ]

    for item, estado, costo, detalle in items_fase_3:
        row_cells = table3.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = estado
        row_cells[2].text = costo
        row_cells[3].text = detalle

    doc.add_paragraph('\n')

    # FASE 4
    doc.add_heading('Fase 4: Compliance, Audit & Operativa', level=1)
    doc.add_paragraph('Estado: EN EJECUCIÓN PARCIAL\nCosto Pendiente Restante: 50 hrs', style='Intense Quote')

    table4 = doc.add_table(rows=1, cols=4)
    table4.style = 'Light Shading Accent 1'
    hdr_cells4 = table4.rows[0].cells
    hdr_cells4[0].text = 'Sub-Fase / Módulo'
    hdr_cells4[1].text = 'Estado'
    hdr_cells4[2].text = 'Costo Estimado'
    hdr_cells4[3].text = 'Detalle Técnico'

    items_fase_4 = [
        ('4.1 CISA/SOX & Tracker Requisitos', 'Completado', '0 hrs', 'Matriz de segregación SoD, y Módulo Gestor de Requerimientos (Tickets) integrados hoy.'),
        ('4.2 Módulo de Alertas Legales', 'Pendiente', '20 hrs', 'Job automático para detectar vencimientos de RNE/RNPA/ISO y envío de correos preventivos.'),
        ('4.3 Real vs Theoretical Loss', 'Pendiente', '30 hrs', 'Analítica para diferenciar Mermas productivas (Pérdidas) de Scrap Vendible (Activo Valorizado).')
    ]

    for item, estado, costo, detalle in items_fase_4:
        row_cells = table4.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = estado
        row_cells[2].text = costo
        row_cells[3].text = detalle


    # Resumen Financiero
    doc.add_heading('Resumen de Costos Estimados (Esfuerzo)', level=2)
    p = doc.add_paragraph()
    p.add_run('Costo Acumulado Fase 1 y entregables de Arquitectura Base: ').bold = True
    p.add_run('Pagado / Consolidado\n')
    p.add_run('Costo Tostalo Proyectado para Fases 2, 3 y 4 (Módulos Finales UI/UX y Motores Asíncronos): ').bold = True
    p.add_run('330 hrs de Ingeniería y Consultoría Técnica.\n\n')
    
    p.add_run('Nota de Ejecución: ').bold = True
    p.add_run('Se recomienda priorizar la "Sub-fase 2.2" (Interfaz Fazón y Liquidación) dado que su arquitectura de Base de Datos y Repositorio Documental ya se encuentra 100% instalada en producción y su retorno operativo es masivo e inmediato para la tercerización de costura/armado.')

    filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Plan_Maestro_Actualizado_MSAC_v4.2.docx')
    doc.save(filepath)
    print(f"Documento Guardado Exitosamente en: {filepath}")

if __name__ == '__main__':
    create_plan_docx()
