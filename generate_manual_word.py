import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm
import os

# Configuración de Matplotlib para no usar GUI
import matplotlib
matplotlib.use('Agg')

DOC_PATH = "docs/MANUAL_SOD_CONTROL_INTERNO.docx"
IMG_DIR = "docs/img"

if not os.path.exists(IMG_DIR):
    os.makedirs(IMG_DIR)

def create_flow_diagram(filename, title, steps):
    """
    Genera un diagrama de flujo vertical limpio con Matplotlib SIN EMOJIS (para compatibilidad).
    steps: lista de tuplas (Rol, Acción)
    """
    fig, ax = plt.subplots(figsize=(8, len(steps) * 1.8 + 1))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, len(steps) * 2.5 + 1)
    ax.axis('off')
    
    # Título del gráfico
    ax.text(5, len(steps) * 2.5 + 0.5, title, ha='center', va='center', fontsize=14, fontweight='bold', color='#333333')
    
    y_pos = len(steps) * 2.5 - 1.5
    box_height = 1.5
    box_width = 8
    x_center = 5
    
    colors = ['#E3F2FD', '#E8F5E9', '#FFF3E0', '#F3E5F5', '#FFEBEE']
    
    for i, (role, action) in enumerate(steps):
        color = colors[i % len(colors)]
        
        # Caja principal
        rect = patches.FancyBboxPatch(
            (x_center - box_width/2, y_pos), box_width, box_height,
            boxstyle="round,pad=0.1", fc=color, ec='#546E7A', lw=2
        )
        ax.add_patch(rect)
        
        # Texto Rol (Izquierda, grande)
        ax.text(x_center - box_width/2 + 0.2, y_pos + box_height - 0.4, role, 
                ha='left', va='center', fontsize=11, fontweight='bold', color='#1565C0')
        
        # Texto Acción (Centro, más pequeño)
        ax.text(x_center, y_pos + box_height/2 - 0.2, action, 
                ha='center', va='center', fontsize=10, wrap=True)
        
        # Flecha hacia abajo (si no es el último)
        if i < len(steps) - 1:
            ax.arrow(x_center, y_pos, 0, -0.6, head_width=0.3, head_length=0.3, fc='#455A64', ec='#455A64')
            y_pos -= 2.5
    
    output_path = os.path.join(IMG_DIR, filename)
    plt.savefig(output_path, bbox_inches='tight', dpi=150)
    plt.close()
    return output_path

def add_audit_matrix(doc):
    doc.add_page_break()
    doc.add_heading('Anexo: Matriz de Riesgos y Controles (Auditoría)', level=1)
    doc.add_paragraph('La siguiente matriz detalla los riesgos operativos identificados y los controles mitigantes implementados mediante la segregación de funciones (SoD).')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Riesgo Identificado'
    hdr_cells[1].text = 'Control de Sistema (SoD)'
    hdr_cells[2].text = 'Roles Separados'
    
    # Datos de la matriz
    data = [
        ("Compras Ficticias / Fraude en Adquisiciones", "Requerimiento de aprobación gerencial y recepción independiente.", "Solicitante vs Aprobador vs Comprador vs Recepcionista"),
        ("Pagos no Autorizados / Malversación", "Three-Way Match (Factura, OC, Recepción) requerido antes del pago.", "Cuentas por Pagar vs Autorizador vs Tesorero"),
        ("Robo de Mercadería / Faltante de Stock", "Separación entre quien factura y quien despacha.", "Facturación vs Almacenista"),
        ("Jineteo de Fondos (Laaping) en Cobranzas", "Separación entre quien emite recibo y quien custodia/deposita valores.", "Cobranzas vs Tesorería vs Contabilidad")
    ]
    
    for riesgo, control, roles in data:
        row_cells = table.add_row().cells
        row_cells[0].text = riesgo
        row_cells[1].text = control
        row_cells[2].text = roles

def generate_word_doc():
    doc = Document()
    
    # Estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Título Principal
    heading = doc.add_heading('Manual de Control Interno y Segregación de Funciones', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    p = doc.add_paragraph('Este documento define el modelo de Control Interno implementado en el sistema MultiMCP, alineado con las mejores prácticas de auditoría y seguridad de la información.')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph('Objetivo: Garantizar la integridad de las operaciones financieras y logísticas mediante una estricta Segregación de Funciones (SoD), asegurando que ninguna persona tenga control total sobre una transacción crítica.')
    doc.add_page_break()

    # --- 1. COMPRAS ---
    doc.add_heading('1. Ciclo de Compras (Procure-to-Pay)', level=1)
    doc.add_paragraph('Control del egreso por adquisiciones. Se mitiga el riesgo de compras innecesarias o a proveedores no autorizados.')
    
    steps_compras = [
        ("SOLICITANTE", "Genera Requerimiento de Compra"),
        ("APROBADOR", "Autoriza la Compra (Control Presupuestario)"),
        ("COMPRADOR", "Gestiona Cotizaciones y Emite Orden de Compra"),
        ("RECEPCIONISTA", "Valida Recepción Física vs Orden de Compra")
    ]
    img_path = create_flow_diagram("flujo_compras_v2.png", "Flujo de Control: Compras", steps_compras)
    doc.add_picture(img_path, width=Inches(6.5))
    doc.add_paragraph('\n')
    doc.add_page_break()

    # --- 2. PAGOS ---
    doc.add_heading('2. Ciclo de Pagos (Treasury Out)', level=1)
    doc.add_paragraph('Control de fondos salientes. Se mitiga el riesgo de pagos duplicados o fraudulentos.')
    
    steps_pagos = [
        ("CUENTAS POR PAGAR", "Registra Factura y Vincula con OC/Remito"),
        ("AUTORIZADOR", "Valida Documentación (3-Way Match) y Autoriza"),
        ("TESORERO", "Ejecuta el Pago (Interbanking / Cheque)"),
        ("CONTADOR", "Concilia Extracto Bancario")
    ]
    img_path = create_flow_diagram("flujo_pagos_v2.png", "Flujo de Control: Pagos", steps_pagos)
    doc.add_picture(img_path, width=Inches(6.5))
    doc.add_paragraph('\n')
    doc.add_page_break()

    # --- 3. VENTAS ---
    doc.add_heading('3. Ciclo de Ventas (Order-to-Cash)', level=1)
    doc.add_paragraph('Control de ingresos y stock. Se mitiga el riesgo de entrega de mercadería sin facturación o ventas a crédito no autorizadas.')
    
    steps_ventas = [
        ("VENDEDOR", "Genera Nota de Pedido"),
        ("FACTURACION", "Verifica Crédito y Emite Factura Fiscal"),
        ("ALMACENISTA", "Prepara Pedido y Despacha Mercadería"),
        ("COBRANZAS", "Gestiona el Cobro de la Factura")
    ]
    img_path = create_flow_diagram("flujo_ventas_v2.png", "Flujo de Control: Ventas", steps_ventas)
    doc.add_picture(img_path, width=Inches(6.5))
    doc.add_paragraph('\n')
    doc.add_page_break()

    # --- 4. COBRANZAS ---
    doc.add_heading('4. Ciclo de Cobranzas (Treasury In)', level=1)
    doc.add_paragraph('Control de recaudación. Se mitiga el riesgo de retención indebida de valores.')
    
    steps_cobranzas = [
        ("COBRANZAS", "Recibe Valor y Emite Recibo Oficial"),
        ("TESORERO", "Custodia Valores y Deposita en Banco"),
        ("CONTADOR", "Audita y Concilia Cuentas Bancarias")
    ]
    img_path = create_flow_diagram("flujo_cobranzas_v2.png", "Flujo de Control: Cobranzas", steps_cobranzas)
    doc.add_picture(img_path, width=Inches(6.5))
    
    # --- ANEXO AUDITORIA ---
    add_audit_matrix(doc)
    
    doc.save(DOC_PATH)
    print(f"Documento generado exitosamente: {DOC_PATH}")

if __name__ == "__main__":
    generate_word_doc()
