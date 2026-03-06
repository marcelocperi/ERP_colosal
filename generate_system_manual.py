import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
import os

# Configuración de Matplotlib
import matplotlib
matplotlib.use('Agg')

DOC_PATH = "docs/MANUAL_GENERAL_SISTEMA.docx"
IMG_DIR = "docs/img"

if not os.path.exists(IMG_DIR):
    os.makedirs(IMG_DIR)

def create_flow_diagram(filename, title, steps):
    """
    Genera un diagrama de flujo vertical limpio con Matplotlib.
    steps: lista de tuplas (Actor/Fase, Acción)
    """
    fig, ax = plt.subplots(figsize=(8, len(steps) * 1.5 + 2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, len(steps) * 2.0 + 1)
    ax.axis('off')
    
    # Título del gráfico
    ax.text(5, len(steps) * 2.0 + 0.5, title, ha='center', va='center', fontsize=14, fontweight='bold', color='#333333')
    
    y_pos = len(steps) * 2.0 - 1.5
    box_height = 1.2
    box_width = 8
    x_center = 5
    
    colors = ['#E3F2FD', '#E8F5E9', '#FFF3E0', '#F3E5F5', '#FFEBEE', '#E0F7FA']
    
    for i, (role, action) in enumerate(steps):
        color = colors[i % len(colors)]
        
        # Caja principal
        rect = patches.FancyBboxPatch(
            (x_center - box_width/2, y_pos), box_width, box_height,
            boxstyle="round,pad=0.1", fc=color, ec='#546E7A', lw=2
        )
        ax.add_patch(rect)
        
        # Texto Rol (Izquierda, grande)
        ax.text(x_center - box_width/2 + 0.2, y_pos + box_height - 0.4, role, 
                ha='left', va='center', fontsize=11, fontweight='bold', color='#1565C0')
        
        # Texto Acción (Centro, más pequeño)
        ax.text(x_center, y_pos + box_height/2 - 0.2, action, 
                ha='center', va='center', fontsize=10, wrap=True)
        
        # Flecha hacia abajo (si no es el último)
        if i < len(steps) - 1:
            ax.arrow(x_center, y_pos, 0, -0.6, head_width=0.3, head_length=0.3, fc='#455A64', ec='#455A64')
            y_pos -= 2.0
    
    output_path = os.path.join(IMG_DIR, filename)
    plt.savefig(output_path, bbox_inches='tight', dpi=150)
    plt.close()
    return output_path

def add_audit_matrix(doc):
    doc.add_page_break()
    doc.add_heading('Anexo: Matriz de Control Interno (Auditoría)', level=1)
    doc.add_paragraph('Matriz de riesgos operativos y controles mitigantes implementados en el sistema.')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Riesgo'
    hdr_cells[1].text = 'Control de Sistema'
    hdr_cells[2].text = 'Mitigación'
    
    data = [
        ("Acceso no autorizado", "Autenticación Multi-Tenant por Sesión", "Aislamiento lógico de empresas"),
        ("Compras Ficticias", "Segregación Solicitante vs Aprobador", "Roles SoD estrictos"),
        ("Pagos Fraudulentos", "Validación 3-Way Match", "Requiere OC y Recepción previas"),
        ("Fuga de Stock", "Trazabilidad de Movimientos", "Registro inmutable de quien despacha")
    ]
    
    for riesgo, control, roles in data:
        row_cells = table.add_row().cells
        row_cells[0].text = riesgo
        row_cells[1].text = control
        row_cells[2].text = roles

def generate_system_manual():
    doc = Document()
    
    # Estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Título Principal
    heading = doc.add_heading('Manual General del Sistema MultiMCP', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Versión 2.0 - Documentación Integral')
    
    doc.add_paragraph('\n')
    doc.add_paragraph('Este manual describe la operatividad completa del sistema, abarcando infraestructura, gestión multi-empresa y módulos operativos. Diseñado para usuarios clave y auditores.')
    doc.add_page_break()

    # --- 1. INFRAESTRUCTURA ---
    doc.add_heading('1. Infraestructura y Arquitectura', level=1)
    doc.add_paragraph('El sistema opera bajo un modelo Multi-Tenant (Múltiples empresas en una sola instalación), garantizando el aislamiento de datos y la seguridad.')
    
    steps_infra = [
        ("USUARIO", "Accede vía Navegador Web (HTTPS)"),
        ("SEGURIDAD", "Firewall y Proxy Inverso (Protección DDOS/SSL)"),
        ("APLICACION", "Servidor Flask (Procesa Lógica de Negocio)"),
        ("DATOS", "Base de Datos MariaDB (Almacenamiento Seguro Aislado)")
    ]
    img_path = create_flow_diagram("infraestructura.png", "Arquitectura del Sistema", steps_infra)
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_page_break()

    # --- 2. GESTIÓN MULTI-EMPRESA ---
    doc.add_heading('2. Gestión de Empresas', level=1)
    doc.add_paragraph('Módulo administrativo para la creación y configuración de nuevos entornos de trabajo.')
    
    steps_empresa = [
        ("ADMIN GLOBAL", "Inicia proceso de Alta de Empresa"),
        ("SISTEMA", "Crea Registro y Copia Datos Maestros (Plan de Cuentas)"),
        ("SEGURIDAD (SoD)", "Genera Automáticamente Roles y Usuarios Base"),
        ("EMPRESA LISTA", "Disponible para operar con seguridad activada")
    ]
    img_path = create_flow_diagram("flujo_empresa.png", "Proceso de Alta de Empresa", steps_empresa)
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_page_break()

    # --- 3. GESTIÓN DE STOCK (INVENTARIO) ---
    doc.add_heading('3. Gestión de Stock e Inventario', level=1)
    doc.add_paragraph('Control total de los bienes de cambio, desde su definición hasta su movimiento físico.')
    
    steps_stock = [
        ("CATALOGADOR", "Define Artículo (Código, Descripción, Tipo)"),
        ("ALMACENISTA", "Registra Entrada (Compra) o Ajuste Inicial"),
        ("SISTEMA", "Actualiza Existencias y Valuación (PPP)"),
        ("AUDITOR", "Realiza Conteos Cíclicos y Ajustes")
    ]
    img_path = create_flow_diagram("flujo_stock.png", "Ciclo de vida del Inventario", steps_stock)
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_page_break()

    # --- 4. COMPRAS (PROCURE-TO-PAY) ---
    doc.add_heading('4. Módulo de Compras e Importaciones', level=1)
    doc.add_paragraph('Ciclo completo de aprovisionamiento con controles de segregación de funciones.')
    
    doc.add_heading('4.1 Identificación y Trazabilidad de Proveedores', level=2)
    doc.add_paragraph('El sistema permite la recuperación ágil de datos de proveedores mediante el uso de Códigos Internos (ej. SUP-0045), Razón Social o CUIT. Esta funcionalidad asegura que cada operación esté correctamente vinculada al maestro de proveedores desde el inicio.')
    
    # Imagen de la nueva interfaz (si existe)
    img_interfaz = os.path.join(os.getcwd(), "nueva_orden_compra_mejorada_1772196785278.png")
    if os.path.exists(img_interfaz):
        doc.add_picture(img_interfaz, width=Inches(6.0))
        doc.add_paragraph('Interfaz de Nueva Orden: Buscador por Código y Panel de validación.')

    steps_compras = [
        ("SOLICITANTE", "Genera Requerimiento (Búsqueda por Código/CUIT)"),
        ("APROBADOR", "Autoriza gasto según presupuesto"),
        ("COMPRADOR", "Gestiona proveedores y emite OC"),
        ("RECEPCIONISTA", "Ingresa mercadería al depósito (3-Way Match)")
    ]
    img_path = create_flow_diagram("flujo_compras.png", "Flujo de Compras Mejorado", steps_compras)
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_page_break()

    # --- 5. VENTAS (ORDER-TO-CASH) ---
    doc.add_heading('5. Módulo de Ventas', level=1)
    doc.add_paragraph('Gestión comercial desde la cotización hasta la facturación.')
    
    steps_ventas = [
        ("VENDEDOR", "Crea Presupuesto y Nota de Pedido"),
        ("FACTURACION", "Emite Comprobante Fiscal (Factura A/B)"),
        ("ALMACENISTA", "Prepara (Picking) y Despacha"),
        ("COBRANZAS", "Gestiona el cobro de la cuenta corriente")
    ]
    img_path = create_flow_diagram("flujo_ventas.png", "Flujo de Ventas", steps_ventas)
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_page_break()

    # --- 6. TESORERÍA (FONDOS) ---
    doc.add_heading('6. Tesorería y Fondos', level=1)
    doc.add_paragraph('Administración centralizada del flujo de dinero (Cash Flow).')
    
    steps_fondos = [
        ("COBRANZAS", "Emite Recibo por cobro a clientes"),
        ("PAGOS", "Emite Orden de Pago a proveedores"),
        ("TESORERO", "Custodia valores y concilia bancos"),
        ("CONTABILIDAD", "Registra asientos contables automáticos")
    ]
    img_path = create_flow_diagram("flujo_fondos.png", "Flujo de Fondos", steps_fondos)
    doc.add_picture(img_path, width=Inches(6.0))
    
    # --- ANEXO ---
    add_audit_matrix(doc)
    
    doc.save(DOC_PATH)
    print(f"Documento generado exitosamente: {DOC_PATH}")

if __name__ == "__main__":
    generate_system_manual()
