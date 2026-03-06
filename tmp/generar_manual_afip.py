
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os
import sys

# Detectar la raíz del proyecto (un nivel arriba de /tmp)
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(project_root)

OUTPUT_PATH = os.path.join(project_root, "tmp", "Manual_AFIP_Security_" + datetime.datetime.now().strftime("%Y%m%d_%H%M%S") + ".docx")

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Crear marcador (bookmark) para el link interno
    bookmark_id = text.replace(' ', '_').replace('.', '').replace('—', '').replace('(', '').replace(')', '').lower()
    add_bookmark(p, bookmark_id)
    return p

def add_bookmark(paragraph, name):
    tag = paragraph._p
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), name)
    tag.insert(0, start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    tag.append(end)

def add_hyperlink(paragraph, text, bookmark_name):
    """
    Agrega un hipervínculo interno a un marcador.
    """
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Estilo de link: azul y subrayado
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    
    text_obj = OxmlElement('w:t')
    text_obj.text = text
    new_run.append(text_obj)
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return paragraph

def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    p.paragraph_format.left_indent = Inches(0.4)
    shading_elm = OxmlElement('w:pPr')
    return p

def separator():
    doc.add_paragraph('─' * 80)

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_picture  # no image needed

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('🚢  NABUCODONOSOR TAX ENGINE')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('Manual Técnico de Facturación Electrónica AFIP/ARCA')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

sub2_p = doc.add_paragraph()
sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = sub2_p.add_run(f'Versión Nebuchadnezzar-v2.1  |  {datetime.date.today().strftime("%d/%m/%Y")}')
run3.font.size = Pt(11)
run3.font.italic = True

doc.add_paragraph()
quote_p = doc.add_paragraph()
quote_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = quote_p.add_run('"No existe una cuchara... pero sí existe el tope de $191.624 para Consumidor Final anónimo."')
r.font.italic = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ═══════════════════════════════════════════════════════════════════════════════
heading('Índice de Contenidos', level=1)
sections = [
    '1. Introducción — La Resistencia vs AFIP',
    '2. Arquitectura del Tax Engine (Nabucodonosor)',
    '3. El Equipo: Actores y sus Roles',
    '4. Controles del Motor (Pre-Flight Check)',
    '5. Gestión de Errores AFIP (Diccionario de Agentes)',
    '6. Protocolo de Resistencia (Reintentos y Cola Zion)',
    '7. Scanner de Traidores (WSAPOC)',
    '8. El Señuelo Scout (FEDummy)',
    '9. Los Inspectores de Identidad (wconsucuit y A10)',
    '10. Vitácora de Vuelo — Registro de Seguridad',
    '11. El Scanner APOC — Detector de Traidores',
    '12. Escudos Activos — Seguridad en Proveedores',
    '13. El Bibliotecario (A100) — Inteligencia de Parámetros',
    '14. Protocolo "La Gran Purga" — Auditoría Profunda',
    '15. Sistema de Tickets de Embarque — Email post-CAE',
    '16. Fragmentos de la Bitácora (Diálogos de Caza)',
    '17. TheKeyMaker + Trinity — Gestión de Tokens',
    '18. Copia Fiel (La Huella Digital)',
    '19. Guía de Activación de Servicios en AFIP',
    '20. Reporte de Campaña Final (Vitácora de Caza)',
]
for s in sections:
    p = doc.add_paragraph(style='No Spacing')
    bookmark_id = s.replace(' ', '_').replace('.', '').replace('—', '').replace('(', '').replace(')', '').lower()
    add_hyperlink(p, s, bookmark_id)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ═══════════════════════════════════════════════════════════════════════════════
heading('1. Introducción — La Resistencia vs AFIP', level=1)
body(
    'El presente manual documenta el diseño, implementación y operación del '
    'Nabucodonosor Tax Engine, un motor de facturación electrónica inteligente '
    'integrado con los Web Services de AFIP/ARCA (WSFE v1, WSAA, WSAPOC).',
    bold=False
)
body(
    'A diferencia de una integración genérica, este motor incorpora múltiples capas de '
    'inteligencia fiscal que permiten: detectar errores localmente antes de contactar a AFIP, '
    'recuperarse automáticamente de caídas del servicio, y notificar a los pasajeros '
    '(clientes) de forma transparente y humana.',
)
body(
    'La metáfora de Matrix fue utilizada durante el desarrollo para nombrar y visualizar '
    'cada componente del sistema, facilitando la comunicación entre el equipo técnico.',
    italic=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ARQUITECTURA
# ═══════════════════════════════════════════════════════════════════════════════
heading('2. Arquitectura del Tax Engine', level=1)
body('El flujo completo de una solicitud de CAE sigue esta secuencia:')
steps = [
    ('PASO 1', 'Pre-Flight Check (local)', 'El Nebuchadnezzar Engine valida matemáticas, fechas, topes y requisitos de servicios SIN contactar a AFIP.'),
    ('PASO 2', 'TheKeyMaker consulta a Trinity', 'Verifica si ya existe un Token de Acceso válido en fin_trinity_tokens. Si existe y no vence en 10 min, lo reutiliza. Si no, fabrica uno nuevo.'),
    ('PASO 3', 'Conexión al Túnel WSFE', 'Abre la conexión SOAP al servicio de Facturación Electrónica de AFIP.'),
    ('PASO 4', 'Envío del FECAESolicitar', 'Envía el lote de comprobantes a AFIP y espera la respuesta.'),
    ('PASO 5', 'Interpretación de Respuesta', 'Si es aprobado: guarda CAE y envía Ticket de Embarque. Si es rechazado por datos: traduce el error. Si es por infraestructura: activa reintentos.'),
    ('PASO 6', 'Protocolo de Resistencia', 'Si los 3 reintentos fallan, guarda en Zion (fin_cae_pendientes) y notifica al usuario con mensaje claro.'),
    ('PASO 7', 'Tank, el Operador', 'Job en background que cada 10 min rescata comprobantes de Zion, reintenta el CAE y envía el email al cliente si lo obtiene.'),
]
for code, title, desc in steps:
    p = doc.add_paragraph()
    r1 = p.add_run(f'  {code} — {title}: ')
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
    p.add_run(desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. EL EQUIPO
# ═══════════════════════════════════════════════════════════════════════════════
heading('3. El Equipo: Actores y sus Roles', level=1)
body('Tabla completa de componentes técnicos y su equivalente en la ficción de Matrix:')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Componente Técnico', 'Nombre en la Nave', 'Actor Matrix', 'Rol en la Ficción', 'Estado']
header_row = table.rows[0]
for i, h in enumerate(headers):
    cell = header_row.cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    set_cell_bg(cell, '1e293b')

rows_data = [
    ('Certificados AFIP (.crt/.key)', 'Las Llaves', '🗝️ Las Llaves', 'Objetos físicos únicos. Sin ellas ninguna puerta se abre. El Keymaker las fabrica una vez y duran años.', '✅ Válidos hasta 2028'),
    ('Cache de Tokens WSAA', 'Trinity (fin_trinity_tokens)', '👩‍💻 Trinity', 'Hackeadora de rasgos orientales. Guarda y reutiliza los tokens. Nunca pide dos veces lo que ya tiene.', '✅ Activa'),
    ('Generador de Tokens WSAA', 'TheKeyMaker (_the_key_maker)', '🧑‍🔬 El Fabricante de Llaves', 'Programa japonés obsesionado con puertas. Fabrica la llave exacta solo cuando la anterior venció. Trabaja en equipo con Trinity.', '✅ Operativo'),
    ('Validaciones Pre-CAE', 'Nebuchadnezzar Engine', '🛸 El Nabucodonosor', 'La nave misma. Chequeo completo antes de cada misión: matemáticas, fechas, topes, servicios. No sale al campo con fallas.', '✅ 4 capas activas'),
    ('Errores 500/503/Timeout', '⚡ Los Sentinels', '🦑 Los Centinelas', 'Máquinas calamar que cortan el túnel. Cuando atacan, la factura no puede pasar. El protocolo de resistencia los enfrenta.', '⚡ Detectados y manejados'),
    ('Errores de datos (10016, IVA)', '🕵️ Los Agentes', '👔 Agente Smith', 'Programas de AFIP que detectan inconsistencias. No perdonan ni un centavo. Se los frena con validación local antes de enviar.', '🛡️ Bloqueados en Pre-vuelo'),
    ('Conexión WSFE/WSAA/WSAPOC', 'Los Túneles', '🌀 Túneles del Keymaker', 'Conexiones encriptadas hacia la Matrix de AFIP. Cada servicio tiene su propio túnel. El sistema detecta cuál cayó.', '✅ WSFE y WSAPOC abiertos'),
    ('Reintento automático en caída', 'Protocolo de Resistencia', '⚔️ La Resistencia', 'Frente a los Sentinels, no huyen: esperan 5s, reagrupan y vuelven al túnel. Máximo 3 intentos antes de retirarse a Zion.', '✅ Activo'),
    ('Cola de facturas sin CAE', 'Zion (fin_cae_pendientes)', '🏙️ Zion', 'El último refugio. Cuando todos los túneles están bloqueados, los datos se guardan seguros hasta poder contraatacar.', '✅ Creada'),
    ('Scanner proveedores falsos', 'El Oráculo APOC', '🔮 El Oráculo', 'Sabe quién es un traidor antes que nadie. Consulta WSAPOC para detectar si un proveedor es un Agente disfrazado (empresa apócrifa).', '⏳ Activar en prod.'),
    ('Email post-CAE al cliente', 'Ticket de Embarque', '🎫 El Pasaporte', 'El documento que el pasajero necesita para demostrar que su viaje es legítimo. Llega por email con el sello de AFIP/ARCA.', '✅ Listo'),
    ('Job de reintento + email', 'Tank, el Operador', '🎮 Tank', 'El único nacido libre a bordo. Monitorea los túneles. Cuando AFIP vuelve, rescata las facturas de Zion y las despacha.', '⏳ Conectar scheduler'),
    ('Proveedor apócrifo en APOC', 'Agente Cypher', '🍖 Cypher', 'El traidor a bordo. Firma un pacto oscuro. En AFIP: empresas que facturan sin capacidad operativa. El Oráculo lo expone.', '🔍 Escaneado por WSAPOC'),
    ('Servicio FEDummy', 'Señuelo Scout', '🛰️ Scout Dummy', 'Un señuelo que se lanza sin credenciales para ver si los túneles están bloqueados por Sentinels antes de arriesgar la nave.', '✅ Operativo'),
    ('Servicio wconsucuit', 'Inspector de Entrada', '🕵️‍♂️ Ghost', 'Tripulante experto en identificación rápida. Valida CUITs a gran velocidad para detectar intrusos en compras y ventas.', '✅ Activo'),
    ('Servicio Padrón A10', 'Scout de Datos', '🦾 Niobe', 'Capitana del Logos. Entra profundo en la Matrix A10 para obtener los datos exactos del intruso y validar la lista de invitados.', '✅ Activa'),
    ('Tabla fin_neb_bitacora', 'Vitácora de Vuelo', '📖 El Libro de Zion', 'Registro persistente de todas las incursiones, vulnerabilidades detectadas y cazas de intrusos. Nada se escapa.', '✅ Grabando'),
]

for row_data in rows_data:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. CONTROLES PRE-FLIGHT
# ═══════════════════════════════════════════════════════════════════════════════
heading('4. Controles del Motor — Pre-Flight Check', level=1)
body('El Nabuchadnezzar Engine ejecuta 4 capas de validación local ANTES de contactar a AFIP. Cada capa es un escudo contra los Agentes:')
doc.add_paragraph()

checks = [
    ('Capa 1 — Topes de Consumidor Final (Anti-Multas)',
     'Verifica si el cliente es Consumidor Final anónimo (sin DNI). Si el importe supera $191.624, bloquea el envío y exige identificación del cliente. Evita rechazo de AFIP y posibles multas por no identificar al comprador.',
     'validar_integridad + chequeo de doc_nro'),
    ('Capa 2 — Integridad Matemática (Sumatoria Cero)',
     'Recalcula: Total esperado = Neto + IVA + Percepciones. Si la diferencia supera $0.01, bloquea el envío con mensaje: "Inconsistencia de Redondeo". AFIP no perdona ni un centavo.',
     'validar_integridad_matematica()'),
    ('Capa 3 — Ventana de Fechas',
     'Valida que la fecha del comprobante no supere el margen permitido: 5 días para Productos (Concepto 1) y 10 días para Servicios (Conceptos 2 y 3). Una factura fuera de ventana es rechazada por AFIP automáticamente.',
     'validar_ventana_fechas()'),
    ('Capa 4 — Requisitos de Servicios',
     'Para Concepto 2 (Servicios) y Concepto 3 (Mixto): verifica que se hayan informado Fecha Desde, Fecha Hasta y Fecha de Vencimiento de Pago del período. Sin estos 3 campos, AFIP rechaza el comprobante.',
     'validar_periodo_servicios()'),
]

for title, desc, method in checks:
    heading(title, level=2)
    body(desc)
    body(f'Método: ', bold=True)
    code_block(f'AfipService.{method}')
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. DICCIONARIO DE ERRORES
# ═══════════════════════════════════════════════════════════════════════════════
heading('5. Diccionario de Errores AFIP (Traducción Humana)', level=1)
body('AFIP devuelve códigos numéricos. El Engine los traduce a mensajes accionables para la tripulación:')
doc.add_paragraph()

err_table = doc.add_table(rows=1, cols=4)
err_table.style = 'Table Grid'
for i, h in enumerate(['Código AFIP', 'Descripción Original', 'Traducción del Engine', 'Acción Sugerida']):
    cell = err_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    set_cell_bg(cell, '4f46e5')

errores = [
    ('10016', 'Error de correlatividad', 'El número local no coincide con el esperado por AFIP. Sincronizar numeración.', 'Ejecutar sincronizar_numeracion()'),
    ('10192', 'FCE requerida', 'Debe emitir Factura de Crédito MiPyME (FCE) por monto y tipo de cliente.', 'Cambiar tipo de comprobante a FCE'),
    ('10015', 'Error en receptor', 'El CUIT/DNI del cliente no es válido en el padrón de AFIP.', 'Verificar datos con consultar_padron()'),
    ('10048', 'No autorizado', 'El CUIT no está autorizado a emitir este tipo de comprobante.', 'Verificar habilitaciones en AFIP'),
    ('500', 'Error interno AFIP', 'Servidor de AFIP con errores internos. Sentinel detectado.', 'Reintento automático (protocolo de resistencia)'),
    ('501', 'Error DB AFIP', 'Error de base de datos en AFIP.', 'Reintento automático'),
    ('502', 'Transacción activa', 'Saturación en servidores AFIP.', 'Esperar y reintentar'),
]
for row_data in errores:
    row = err_table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. PROTOCOLO DE RESISTENCIA
# ═══════════════════════════════════════════════════════════════════════════════
heading('6. Protocolo de Resistencia — Reintentos y Cola Zion', level=1)

body('Cuando los Sentinels (errores de infraestructura de AFIP) atacan, el sistema ejecuta este protocolo:')
doc.add_paragraph()

protocol = [
    ('Fase 1', 'Clasificación del error', 'El Engine distingue entre errores de infraestructura (500, 503, timeout, SOAP) y errores de negocio (10016, 10015). Los errores de negocio no se reintentan porque AFIP no cambiará de opinión.'),
    ('Fase 2', 'Backoff silencioso', '3 reintentos automáticos con 5 segundos de pausa entre cada uno. El usuario no ve nada, la operación continúa en segundo plano.'),
    ('Fase 3', 'Repliegue a Zion', 'Si los 3 reintentos fallan, el comprobante se guarda en fin_cae_pendientes con estado PENDIENTE. La factura puede imprimirse sin CAE y regularizarse durante el día.'),
    ('Fase 4', 'Notificación Transparente', 'El cajero ve un mensaje claro: "AFIP no responde luego de 3 intentos. La factura fue guardada en estado PENDIENTE DE CAE. El sistema reintentará automáticamente en 10 minutos."'),
    ('Fase 5', 'Tank rescata desde Zion', 'El job procesar_cae_pendientes corre cada 10 minutos. Para cada comprobante pendiente, reintenta el CAE. Si tiene éxito, actualiza la BD y envía el Ticket de Embarque por email al cliente.'),
]

for fase, title, desc in protocol:
    heading(f'{fase}: {title}', level=2)
    body(desc)

heading('Tabla fin_cae_pendientes', level=2)
body('Estructura de la cola de comprobantes sin CAE:')
code_block("""CREATE TABLE fin_cae_pendientes (
    id              INT AUTO_INCREMENT PRIMARY KEY,
    enterprise_id   INT NOT NULL,
    comprobante_id  INT NOT NULL,
    intentos        INT DEFAULT 0,
    ultimo_intento  DATETIME,
    proximo_intento DATETIME,
    ultimo_error    TEXT,
    estado          VARCHAR(20) DEFAULT 'PENDIENTE',
    creado_en       DATETIME DEFAULT CURRENT_TIMESTAMP
);""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. WSAPOC
# ═══════════════════════════════════════════════════════════════════════════════
heading('7. Scanner de Traidores — WSAPOC', level=1)
body('El servicio WSAPOC de AFIP/ARCA permite consultar si un CUIT figura en la Base de Contribuyentes Apócrifos.')
body('Un contribuyente apócrifo es una empresa que emite facturas sin tener capacidad operativa real (empresa fantasma o Agente encubierto en términos de Matrix).')

heading('¿Por qué es crítico?', level=2)
body('Si una empresa acepta facturas de proveedores apócrifos como crédito fiscal, AFIP puede desconocer ese crédito y generar ajustes fiscales significativos. El Scanner detecta esto ANTES de que el daño ocurra.')

heading('Activación del Servicio', level=2)
steps_apoc = [
    'Ingresar con Clave Fiscal al portal de AFIP/ARCA.',
    'Ir a "Administrador de Relaciones de Clave Fiscal".',
    'Seleccionar la empresa representada.',
    'Clic en "Adherir Servicio" → AFIP → Web Services.',
    'Buscar y seleccionar: "Consulta contribuyentes apócrifos (wsapoc)".',
    'Asociar el certificado TESTCOLOSAL (el mismo que usa wsfe).',
]
for i, step in enumerate(steps_apoc, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(step)

heading('Método en el Engine', level=2)
code_block('resultado = await AfipService.consultar_base_apoc(enterprise_id, cuit_proveedor)')
body('Respuesta posible:')
code_block("""# Si está limpio:
{"success": True, "es_apocrifo": False, "mensaje": "✅ CUIT no registra antecedentes apócrifos."}

# Si es un traidor:
{"success": True, "es_apocrifo": True, "mensaje": "⚠️ ALERTA ROJA: CUIT está en la Base APOC de AFIP.", "detalles": "..."}""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. EL SEÑUELO SCOUT
# ═══════════════════════════════════════════════════════════════════════════════
heading('8. El Señuelo Scout — FEDummy', level=1)
body('Antes de iniciar cualquier misión crítica o chequeo de salud, el Operador lanza un Señuelo Scout (FEDummy).')
body('Este servicio de AFIP permite verificar el estado de los túneles sin necesidad de certificados ni autenticación. '
     'Si el señuelo no regresa, sabemos que los Sentinels han tomado el control de la red.')

heading('Diagnóstico de Túneles', level=2)
body('El Scout devuelve el estado de tres motores clave de AFIP:')
items_dummy = [
    'AppServer: El servidor de aplicaciones que recibe las facturas.',
    'DbServer: El motor de base de datos donde se graban los CAEs.',
    'AuthServer: El servidor de autenticación de AFIP.'
]
for item in items_dummy:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

heading('Integración en Health Check', level=2)
body('El método health_check() ahora lanza automáticamente el señuelo antes de reportar el estado de la nave.')
code_block("""# Resultado del escaneo del señuelo:
{
    "success": True,
    "app_server": "OK",
    "db_server": "OK",
    "auth_server": "OK",
    "mensaje": "Túneles escaneados: AppServer=OK, DbServer=OK, AuthServer=OK"
}""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. INSPECTORES DE IDENTIDAD
# ═══════════════════════════════════════════════════════════════════════════════
heading('9. Los Inspectores de Identidad — wconsucuit y A10', level=1)
body('Para evitar que intrusos se infiltren en nuestras operaciones comerciales, hemos sumado dos nuevos tripulantes especialistas:')

heading('wconsucuit (Identificación Rápida)', level=2)
body('Utiliza el Padrón A13 para una validación veloz. Es ideal para compras y ventas donde necesitamos saber quién es el sujeto en milisegundos.')
code_block('await AfipService.consultar_cuit(enterprise_id, cuit)')

heading('Padrón A10 (Scout de Datos)', level=2)
body('Cuando la identificación rápida no es suficiente, Niobe (A10) entra a buscar el historial completo. Valida si lo que tenemos en nuestra lista de invitados es correcto.')
code_block('await AfipService.consultar_datos_a10(enterprise_id, cuit)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. VITÁCORA DE VUELO
# ═══════════════════════════════════════════════════════════════════════════════
heading('10. Vitácora de Vuelo — Registro de Seguridad', level=1)
body('No nos guardamos secretos. Cada excursión de caza de intrusos y cada vulnerabilidad identificada queda grabada en la vitácora.')

heading('Estructura de la Bitácora', level=2)
code_block("""CREATE TABLE fin_neb_bitacora (
    id            INT AUTO_INCREMENT PRIMARY KEY,
    enterprise_id INT NOT NULL,
    evento        VARCHAR(100) NOT NULL,
    tipo          VARCHAR(50) DEFAULT 'INFO', -- INFO, ALERT, SECURITY, ERROR
    detalle       TEXT,
    data_json     JSON,
    creado_en     DATETIME DEFAULT CURRENT_TIMESTAMP
);""")

heading('Eventos de Caza', level=2)
body('Tipos de eventos registrados:')
items_log = [
    'ESCANE_CUIT: Cuando wconsucuit o A10 identifican un sujeto.',
    'VULNERABILIDAD_IDENTIFICADA: Cuando se detectan fallos en los túneles o servicios.',
    'AUDITORIA_INVITADOS: Comparativa entre datos locales y datos de la Matrix.'
]
for item in items_log:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. SCANNER APOC
# ═══════════════════════════════════════════════════════════════════════════════
heading('11. El Scanner APOC — Detector de Traidores', level=1)
body('En la Matrix, un Traidor (Contribuyente Apócrifo) puede comprometer toda la red fiscal de la nave. '
     'El Scanner APOC realiza incursiones profundas en los listados de AFIP para identificar sujetos marcados por irregularidades.')

heading('Protocolo de Detección', level=2)
body('El scanner no solo busca un CUIT; analiza el veredicto de AFIP y levanta escudos inmediatos si detecta una anomalía.')
code_block('await AfipService.consultar_base_apoc(enterprise_id, cuit)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 12. ESCUDOS ACTIVOS: MÓDULO DE COMPRAS
# ═══════════════════════════════════════════════════════════════════════════════
heading('12. Escudos Activos — Seguridad en Proveedores', level=1)
body('La tripulación de AFIP (Scouts, Inspectores y Scanners) ha sido desplegada en el frente de batalla: el Módulo de Compras.')

heading('Check-in de Nuevos Proveedores', level=2)
body('Cada vez que un nuevo sujeto intenta abordar la nave (registro de proveedor), se ejecutan los siguientes protocolos:')
items_compras = [
    'Barrera APOC: Si el sujeto es un Traidor, el acceso es denegado preventivamente.',
    'Validación Niobe: Se cruza el nombre ingresado con el registro oficial de AFIP (A10).',
    'Registro en Vitácora: Cada registro queda auditado con su veredicto de seguridad.'
]
for item in items_compras:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

heading('Bloqueo de Pagos Preventivo', level=2)
body('Incluso si un sujeto logró infiltrarse en el pasado, el Scanner APOC actúa como una última línea de defensa '
     'antes de procesar cualquier Orden de Pago. Si el proveedor ha sido comprometido recientemente, el pago es bloqueado.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 13. EL BIBLIOTECARIO (A100)
# ═══════════════════════════════════════════════════════════════════════════════
heading('13. El Bibliotecario — ws_sr_padron_a100', level=1)
body('El Bibliotecario no es un soldado; es el guardián de la sabiduría de la Matrix AFIP. '
     'Su función es descargar las Tablas Maestras de parámetros para asegurar que la nave hable el mismo idioma que el Sistema Registral.')

heading('El Uso Estratégico', level=2)
body('Evita la corrupción de datos mediante la sincronización oficial de:')
items_a100 = [
    'Provincias y Códigos Postales: El mapeo exacto de la Matrix.',
    'Tipos de Documento e IVA: Las reglas del juego fiscal.',
    'Dependencias de AFIP: El mapa de las oficinas de control.'
]
for item in items_a100:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

code_block('await AfipService.consultar_parametros_a100(enterprise_id, "provincias")')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 14. LA GRAN PURGA
# ═══════════════════════════════════════════════════════════════════════════════
heading('14. Protocolo "La Gran Purga" — Auditoría Profunda', level=1)
body('Incluso una nave blindada puede tener parásitos que se ocultaron antes de activar los escudos. '
     'La Gran Purga es un protocolo de auditoría general que escanea toda la base de datos de proveedores y clientes.')

heading('Acciones de Auditoría', level=2)
code_block('await AfipService.ejecutar_auditoria_general(enterprise_id)')
body('Durante la purga, el sistema realiza:')
items_purge = [
    'Escaneo APOC Masivo: Identifica traidores ocultos entre los invitados.',
    'Cruce de Datos Niobe (A10): Detecta discrepancias entre el nombre cargado y el registro oficial.',
    'Inyección en Vitácora: Cada anomalía descubierta es documentada para su purga posterior.'
]
for item in items_purge:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 15. TICKETS DE EMBARQUE
# ═══════════════════════════════════════════════════════════════════════════════
heading('15. Sistema de Tickets de Embarque — Email post-CAE', level=1)
body('Una vez que el CAE es obtenido exitosamente (ya sea en tiempo real o tras reintentos del Operador), el sistema envía automáticamente el Ticket de Embarque al cliente por email.')

heading('Contenido del email', level=2)
body('El email incluye:')
items = ['Tipo y número de comprobante (ej: Factura B 00001-00000066)', 'Fecha de emisión', 'Importe total', 'Número de CAE otorgado por AFIP', 'Fecha de vencimiento del CAE', 'Sello de validez: "Este comprobante ha sido informado a AFIP/ARCA correctamente"']
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

heading('Activación', level=2)
code_block("""from services.email_service import enviar_ticket_embarque

ok, err = enviar_ticket_embarque(
    destinatario_email = 'cliente@email.com',
    destinatario_nombre = 'Juan Pérez',
    comprobante_data = {
        'tipo_nombre': 'Factura B',
        'punto_venta': 1,
        'numero': 66,
        'fecha_emision': '2026-03-05',
        'total': 15000.00,
        'cae': '74123456789012',
        'cae_vto': '20260315',
    },
    enterprise_id = 0
)""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 18. TRINITY + KEYMAKER
# ═══════════════════════════════════════════════════════════════════════════════
heading('18. TheKeyMaker + Trinity — Gestión Inteligente de Tokens', level=1)

body('El sistema de autenticación con AFIP (WSAA) ha sido mejorado con un sistema de caché inteligente para evitar el error "El CEE ya posee un TA válido".')

heading('El problema original', level=2)
body('AFIP emite un Ticket de Acceso (TA) que tiene una validez de 12 horas. Si el sistema pedía un nuevo ticket mientras el anterior seguía vigente, AFIP rechazaba el request con este error. En producción, con múltiples facturas simultáneas, esto causaba fallos en cadena.')

heading('La solución: Trinity guarda, TheKeyMaker fabrica', level=2)
body('La tabla fin_trinity_tokens actúa como bóveda de llaves. El flujo es:')
flow = [
    ('TheKeyMaker consulta a Trinity', 'Busca en fin_trinity_tokens si hay un token vigente (que no venza en los próximos 10 min).'),
    ('Si Trinity tiene la llave', 'La devuelve inmediatamente. No se contacta a AFIP. Tiempo de respuesta: milisegundos.'),
    ('Si Trinity no tiene la llave', 'TheKeyMaker fabrica un nuevo token firmando el TRA con el certificado, llama al WSAA de AFIP y guarda el resultado en Trinity con su fecha de expiración.'),
    ('UPSERT inteligente', 'Si ya existe un registro para esa empresa+servicio, lo actualiza. Nunca hay duplicados.'),
]
for title, desc in flow:
    p = doc.add_paragraph()
    r = p.add_run(f'  {title}: ')
    r.bold = True
    r.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
    p.add_run(desc)

heading('Tabla fin_trinity_tokens', level=2)
code_block("""CREATE TABLE fin_trinity_tokens (
    id            INT AUTO_INCREMENT PRIMARY KEY,
    enterprise_id INT NOT NULL,
    servicio      VARCHAR(50) NOT NULL,   -- wsfe, wsapoc, ws_sr_padron_a13, etc.
    token         TEXT NOT NULL,
    sign          TEXT NOT NULL,
    expira_en     DATETIME NOT NULL,
    creado_en     DATETIME DEFAULT CURRENT_TIMESTAMP,
    UNIQUE KEY uniq_ent_srv (enterprise_id, servicio)
);""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. GUÍA DE ACTIVACIÓN
# ═══════════════════════════════════════════════════════════════════════════════
heading('10. Guía de Activación de Servicios en AFIP', level=1)
body('Para que el Nabucodonosor pueda abrir los túneles, el CUIT debe tener los siguientes servicios asociados al certificado en el portal de AFIP:')

srvs = [
    ('wsfe', 'Facturación Electrónica (CAE)', 'OBLIGATORIO', 'Motor principal de emisión de facturas.'),
    ('wsapoc', 'Consulta Apócrifos', 'RECOMENDADO', 'Scanner de proveedores fraudulentos.'),
    ('wsremarba', 'Remito Electrónico ARBA (COT)', 'OPCIONAL', 'Para autorización de remitos en Provincia de Buenos Aires.'),
    ('oconws', 'Trazabilidad (Precursores/Farmacia)', 'OPCIONAL', 'Solo aplica para rubros regulados.'),
    ('ws_sr_padron_a13', 'Padrón de Contribuyentes A13', 'RECOMENDADO', 'Consulta de datos de clientes/proveedores por CUIT.'),
]

srv_table = doc.add_table(rows=1, cols=4)
srv_table.style = 'Table Grid'
for i, h in enumerate(['Servicio', 'Nombre', 'Prioridad', 'Uso']):
    cell = srv_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    set_cell_bg(cell, '1e293b')
for s in srvs:
    row = srv_table.add_row()
    for i, val in enumerate(s):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        if i == 2 and val == 'OBLIGATORIO':
            row.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
            row.cells[i].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. REPORTE FINAL
# ═══════════════════════════════════════════════════════════════════════════════
heading('13. Reporte de Campaña Final — Vitácora de Caza', level=1)
body(f'Fecha de cierre de campaña: {datetime.datetime.now().strftime("%d/%m/%Y %H:%M")}', italic=True)
doc.add_paragraph()
body('¡La primera campaña del Nabucodonosor ha sido exitosa! Los logros alcanzados son:', bold=True)

logros = [
    '✅ Certificados validados y vigentes hasta 2028.',
    '✅ Túnel WSFE abierto y operativo.',
    '✅ Túnel WSAPOC abierto (requiere activación en producción).',
    '✅ Trinity activa: tokens cacheados en fin_trinity_tokens.',
    '✅ TheKeyMaker operativo: no más errores de "ya posee TA".',
    '✅ 4 capas de Pre-vuelo implementadas (matemáticas, fechas, topes, servicios).',
    '✅ Diccionario de errores AFIP con traducción humana.',
    '✅ Protocolo de Resistencia: 3 reintentos + backoff.',
    '✅ Cola Zion (fin_cae_pendientes) creada.',
    '✅ Jet de Tickets de Embarque (email post-CAE) listo.',
    '✅ Scanner de Traidores WSAPOC implementado.',
    '✅ Señuelo Scout (FEDummy) integrado en Health Check.',
    '⏳ Pendiente: Conectar Tank (procesar_cae_pendientes) al APScheduler.',
    '⏳ Pendiente: Activar wsapoc en portal de AFIP producción.',
]
for l in logros:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(l)

doc.add_paragraph()
separator()
doc.add_paragraph()

final_p = doc.add_paragraph()
final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = final_p.add_run('"Bienvenido al mundo real. El Nabucodonosor ya factura en él."')
r.font.italic = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)

# ── Guardar ───────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"✅ Manual generado exitosamente en:\n{OUTPUT_PATH}")
