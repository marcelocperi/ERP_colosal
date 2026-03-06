
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_detailed_manual():
    doc = Document()
    
    # Estilo de Títulos
    title = doc.add_heading('Manual de Gobierno de Segregación de Funciones (SoD) y Auditoría Superior', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Sistema: Colosal ERP / BibliotecaWEB\nVersión: 2026.1\nMarco Normante: CISA, SOX y Normas de Auditoría Local').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 1. Generación de Módulos (Gobernanza) ---
    doc.add_heading('1. Ciclo de Vida del Módulo (Generación y Registro)', level=1)
    doc.add_paragraph(
        'Los módulos en Colosal ERP no son solo vistas, son unidades de control auditables. '
        'Para generar un módulo, se siguen estos pasos:'
    )
    doc.add_paragraph('A. Declaración Técnica: Se crea el blueprint en Flask (core/routes.py) definiendo el endpoint funcional.', style='List Bullet')
    doc.add_paragraph('B. Registro en el Diccionario de Permisos: Se debe inscribir un código único en la tabla `sys_permissions` (ej: `create_pago`).', style='List Bullet')
    doc.add_paragraph('C. Persistencia en Menú: Se integra en `.agent/menu_structure.json`, vinculando la ruta con el permiso requerido. '
                    'Este desacoplamiento garantiza que la navegación sea gobernada por el sistema de seguridad.', style='List Bullet')

    # --- 2. Inscripción en Roles ---
    doc.add_heading('2. Inscripción de Permisos en Roles', level=1)
    doc.add_paragraph(
        'La asignación de permisos se realiza a través del panel de Administración de Roles. '
        'Cuando un usuario selecciona permisos para un rol, el sistema ejecuta una validación cruzada '
        'con la matriz SoD definida en `services/sod_service.py`.'
    )

    # --- 3. Funcionamiento del Control SoD ---
    doc.add_heading('3. Motor de Control de Segregación de Funciones (SoD)', level=1)
    doc.add_paragraph(
        'El control SoD (Segregation of Duties) funciona mediante Clusters Funcionales. '
        'Un cluster agrupa permisos que representan una fase de un proceso de negocio.'
    )
    doc.add_paragraph('Clusters Críticos Implementados:', style='List Bullet')
    doc.add_paragraph('– Cluster Compras: `create_orden_compra`, `admin_proveedores`.', style='List Bullet')
    doc.add_paragraph('– Cluster Pagos: `create_pago`, `admin_medios_pago`.', style='List Bullet')
    doc.add_paragraph('– Cluster Stock: `receive_stock`, `admin_depositos`.', style='List Bullet')
    doc.add_heading('Detección de Violaciones:', level=2)
    doc.add_paragraph(
        'El motor `analyze_role_sod` realiza una operación de conjuntos (sets) comparando los permisos '
        'asignados contra las reglas de conflicto. Si un rol tiene permisos de dos clusters antagónicos, '
        'se dispara una Alerta de Violación.'
    )

    # --- 4. Mitigación y Revocación ---
    doc.add_heading('4. Mitigación de Riesgos y Revocación', level=1)
    doc.add_paragraph(
        'Cuando se detecta un riesgo, el auditor puede proceder de dos formas:'
    )
    doc.add_paragraph('1. Mitigación de Auditoría: Documentar por qué el riesgo es aceptable bajo controles compensatorios.', style='List Number')
    doc.add_paragraph('2. Revocación Inmediata: Utilizar el botón de eliminación en el modal de SoD. '
                    'Esta acción remueve la entrada de la tabla `sys_role_permissions` y genera un log de REVOKE_PERMISSION.', style='List Number')

    # --- 5. Registro de Logs (Trazabilidad CISA) ---
    doc.add_heading('5. Registro en Tablas de Logs (Audit Trail)', level=1)
    doc.add_paragraph(
        'Cada acción de gobierno sobre los roles genera una entrada en `sys_security_logs` con:'
    )
    doc.add_paragraph('– dt_date_update: Timestamp exacto del servidor.', style='List Bullet')
    doc.add_paragraph('– Violaciones: Detalle de las reglas SoD que el rol está vulnerando al momento de guardarse.', style='List Bullet')
    doc.add_paragraph('– Inocuos: Conteo de permisos que no representan riesgo transaccional.', style='List Bullet')
    doc.add_paragraph('– Operador e IP: Identificación única del responsable administrativo.', style='List Bullet')

    # --- 6. Auditoría de Campo (Abuso de Perfiles) ---
    doc.add_heading('6. Consulta de Abuso de Perfiles en el Negocio', level=1)
    doc.add_paragraph(
        'Para auditar transacciones pasadas cometidas con roles excesivos, el sistema ofrece el módulo '
        'de Integridad Transaccional.'
    )
    doc.add_heading('Detección Retroactiva:', level=2)
    doc.add_paragraph(
        'El sistema cruza dinámicamente el `user_id` en las tablas de negocio (ej: `fin_ordenes_pago`) '
        'con el análisis SoD del rol que tenía asignado ese usuario. Si un pago fue hecho por un usuario '
        'que también podía pedir compras, se etiqueta como "Operación con Violación".'
    )

    doc.add_heading('Conclusión de Robustez', level=1)
    doc.add_paragraph(
        'El esquema Colosal garantiza un control de 360 grados: desde el diseño del código (módulos), '
        'pasando por la configuración (roles), hasta el monitoreo del negocio (auditoría de integridad).'
    )

    # Guardar
    output_path = r'C:\Users\marce\Documents\GitHub\bibliotecaweb\audit_compliance\Manual_Gobierno_SoD_Avanzado.docx'
    doc.save(output_path)
    print(f"Manual detallado generado en: {output_path}")

if __name__ == "__main__":
    generate_detailed_manual()
