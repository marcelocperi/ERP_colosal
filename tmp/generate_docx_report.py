
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_report():
    doc = Document()

    # Título Principal
    title = doc.add_heading('Informe de Cumplimiento: Marco de Control Interno y Auditoría de Sistemas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadatos del Documento
    p = doc.add_paragraph()
    p.add_run('Referencia: ').bold = True
    p.add_run('Estándares CISA (ISACA), SOX (Sección 404) y Regulación Contable Local.\n')
    p.add_run('Sistema: ').bold = True
    p.add_run('Colosal ERP / BibliotecaWEB Multi-Tenant\n')
    p.add_run('Fecha: ').bold = True
    p.add_run('27 de febrero de 2026')

    # Sección 1
    doc.add_heading('1. Introducción y Propósito', level=1)
    doc.add_paragraph(
        'El presente documento certifica la robustez de los controles implementados en el sistema Colosal ERP, '
        'diseñados bajo los lineamientos del Certified Information Systems Auditor (CISA) y la Ley Sarbanes-Oxley (SOX). '
        'La arquitectura de seguridad se basa en el principio de Defensa en Profundidad, garantizando que ningún '
        'individuo posea control total sobre una transacción financiera de extremo a extremo.'
    )

    # Sección 2
    doc.add_heading('2. Marco de Segregación de Funciones (SoD)', level=1)
    doc.add_paragraph(
        'Colosal implementa un Motor de Análisis Capilar de Permisos (sod_service.py). '
        'Este motor evalúa la matriz de acceso en tiempo real basándose en Clusters Funcionales.'
    )
    
    doc.add_heading('2.1 Reglas de Conflicto Crítico', level=2)
    doc.add_paragraph(
        'Se han definido intersecciones prohibidas entre clusters de Compras, Pagos, Recepción y Contabilidad. '
        'El sistema utiliza un algoritmo que impide que un Comprador sea también Pagador (Violación SOX 404 sobre integridad de desembolsos).'
    )

    # Sección 3
    doc.add_heading('3. Integridad y Trazabilidad (Audit Trail)', level=1)
    doc.add_paragraph(
        'Se ha implementado un esquema de trazabilidad de alto nivel que cumple con el estándar de No Repudio.'
    )
    
    # Tabla de Controles
    doc.add_heading('Tabla de Implementación Técnica de Controles', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Control'
    hdr_cells[1].text = 'Descripción CISA'
    hdr_cells[2].text = 'Implementación en Colosal'
    
    data = [
        ('Preventivo', 'Prevenir errores/fraudes.', 'Bloqueo de asignación masiva de permisos conflictivos.'),
        ('Detectivo', 'Identificar transacciones sospechosas.', 'Reporte de Integridad Transaccional de campo.'),
        ('Correctivo', 'Mitigar riesgos detectados.', 'Módulo de revocación inmediata de permisos excesivos.')
    ]
    
    for ctrl, desc, impl in data:
        row_cells = table.add_row().cells
        row_cells[0].text = ctrl
        row_cells[1].text = desc
        row_cells[2].text = impl

    # Sección 4
    doc.add_heading('4. Conclusión para el Auditor', level=1)
    doc.add_paragraph(
        'La arquitectura de Colosal ERP no es solo una base de datos; es un ecosistema de control auditado que: '
        '1) Detecta anomalías, 2) Documenta alteraciones de privilegios, 3) Certifica integridad funcional.'
    )
    
    p_final = doc.add_paragraph()
    p_final.add_run('Certificado de Robustez: ').bold = True
    p_final.add_run('La implementación de SoD y Auditoría de Campo cumple satisfactoriamente con los criterios de auditoría internacional.')

    # Guardar en Directorio de Artefactos
    output_path = r'C:\Users\marce\.gemini\antigravity\brain\81f2a74f-c624-48a5-9597-bbe69dc58da0\Informe_Cumplimiento_CISA_SOX.docx'
    doc.save(output_path)
    print(f"Documento generado exitosamente en: {output_path}")

if __name__ == "__main__":
    generate_report()
