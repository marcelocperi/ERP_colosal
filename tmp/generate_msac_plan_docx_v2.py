
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill_color):
    """Sets the background color of a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Estilo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    table_data = []
    in_table = False

    for line in lines:
        line = line.strip()
        
        # Saltos de línea vacíos
        if not line:
            if not in_table:
                doc.add_paragraph()
            continue

        # Headings
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            continue
        
        # Tablas (Markdown format)
        if line.startswith('|'):
            if '---' in line:
                continue
            # Clean empty strings from split if they are just separators
            cells = [c.strip() for c in line[1:-1].split('|')]
            
            if not in_table:
                in_table = True
                table_data = [cells]
            else:
                table_data.append(cells)
            continue
        else:
            if in_table:
                # Flush table
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_data):
                    for j, val in enumerate(row_data):
                        if j < len(table.columns):
                            cell = table.cell(i, j)
                            cell.text = val
                            if i == 0: # Header
                                set_cell_background(cell, "D9D9D9")
                                if cell.paragraphs:
                                    cell.paragraphs[0].runs[0].bold = True
                in_table = False
                table_data = []

        # Alerts / Blockquotes
        if line.startswith('>'):
            p = doc.add_paragraph()
            run = p.add_run(line[1:].strip())
            run.italic = True
            font = run.font
            font.color.rgb = RGBColor(0, 102, 204)
            continue

        # Texto normal con posible negrita
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)

    if in_table:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_data):
            for j, val in enumerate(row_data):
                if j < len(table.columns):
                    cell = table.cell(i, j)
                    cell.text = val
                    if i == 0:
                        set_cell_background(cell, "D9D9D9")
                        if cell.paragraphs:
                             cell.paragraphs[0].runs[0].bold = True

    doc.save(docx_path)
    print(f"Documento guardado en: {docx_path}")

if __name__ == "__main__":
    md_file = r'c:\Users\marce\Documents\GitHub\bibliotecaweb\multiMCP\master_plan_msac.md'
    docx_file = r'c:\Users\marce\Documents\GitHub\bibliotecaweb\multiMCP\Plan_Maestro_MSAC_v2.docx'
    convert_md_to_docx(md_file, docx_file)
