import sys
import os

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx no está instalado. Ejecute: pip install python-docx")
    sys.exit(1)

def extract_content(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        return f.readlines()

def create_docx(md_lines, output_path):
    doc = Document()
    
    # Estilos del documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Titulo principal
    title = doc.add_heading('Plan Maestro de Costos Industriales, Sourcing & R&D (MSAC v4.1)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    current_table = None

    for line in md_lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith('# '):
            continue # Ya pusimos el titulo
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('> [!'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.color.rgb = RGBColor(0, 102, 204) # Azul
        elif line.startswith('|'):
            # Handling tables basic parsing
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if '---' in line:
                continue # Skip separator
            
            if current_table is None:
                current_table = doc.add_table(rows=1, cols=len(cells))
                current_table.style = 'Table Grid'
                hdr_cells = current_table.rows[0].cells
                for i, text in enumerate(cells):
                    hdr_cells[i].text = text.replace('*', '')
            else:
                row_cells = current_table.add_row().cells
                for i, text in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = text.replace('*', '')
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            current_table = None # Si habia tabla, se corto
            doc.add_paragraph(line.replace('**', ''))

    doc.save(output_path)
    print(f"Documento generado exitosamente en: {output_path}")

if __name__ == '__main__':
    md_file = 'master_plan_msac.md'
    docx_file = 'Master_Plan_MSAC_v4_1.docx'
    
    if os.path.exists(md_file):
        lines = extract_content(md_file)
        create_docx(lines, docx_file)
    else:
        print(f"No se encontro el archivo {md_file}")
