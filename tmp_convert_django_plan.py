import os
from docx import Document
from docx.shared import Pt

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"File {md_path} not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Estilo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in lines:
        line = line.strip()
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- [ ]'):
            doc.add_paragraph('☐ ' + line[5:].strip())
        elif line.startswith('- [x]') or line.startswith('- [X]'):
            doc.add_paragraph('☑ ' + line[5:].strip())
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif not line:
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Successfully converted to {docx_path}")

if __name__ == "__main__":
    md_file = r"C:\Users\marce\Documents\GitHub\Colosal\PLAN_DE_MIGRACION_DJANGO.md"
    docx_file = r"C:\Users\marce\Documents\GitHub\Colosal\PLAN_DE_MIGRACION_DJANGO.docx"
    md_to_docx(md_file, docx_file)
