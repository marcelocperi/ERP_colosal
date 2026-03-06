import docx
import sys

def append_to_word(filepath):
    try:
        doc = docx.Document(filepath)
    except Exception as e:
        print(f"Error open: {e}")
        return
    
    # Add new section
    doc.add_page_break()
    doc.add_heading('Anexo: Auditoría de Permisos vs. Rutas y Nuevos Módulos Integrados', level=1)
    
    doc.add_paragraph(
        'Como parte del proceso de aseguramiento de la plataforma y soporte al framework FMECA, '
        'se realizó una auditoría exhaustiva sobre todas las rutas registradas en la aplicación (routes.py de todos los módulos) '
        'para verificar su consistencia frente al repositorio maestro de permisos (la tabla sys_permissions).'
    )
    
    doc.add_heading('1. Resultado de la Auditoría de Permisos', level=2)
    p = doc.add_paragraph('El resultado del recorrido por todo el código fuente del servidor (183 rutas mapeadas) arrojó los siguientes resultados:')
    
    ul_style = 'List Bullet'
    doc.add_paragraph('Permisos Faltantes (Missing): 0 rutas. Todos los permisos solicitados explícitamente por el código fuente se encuentran dados de alta correctamente en la base de datos.', style=ul_style)
    doc.add_paragraph('Permisos Confirmados (OK): 43 rutas que están correctamente configuradas bajo el decorador @permission_required(...) mapeadas uno-a-uno con el maestro de configuraciones y distribuidas por dominio (Stock, Compras, Core, Ventas, Fondos, Biblioteca).', style=ul_style)
    doc.add_paragraph('Rutas que requieren exclusivamente sesión activa (@login_required): 130 rutas. Son mayoritariamente endpoints de APIs internas (georef, crons), endpoints de consulta del dashboard global o auxiliares que no dictan operaciones de alto riesgo per se o están implícitamente controladas dentro del flujo posterior.', style=ul_style)
    doc.add_paragraph('Rutas Públicas: 10 rutas (login, request password, reseteo, crons automáticos).', style=ul_style)
    
    doc.add_heading('2. Nuevos Pantallas y Configuración de Perfiles Duales (Templates Añadidos)', level=2)
    doc.add_paragraph(
        'El sistema ha sido enriquecido con tres nuevos templates HTML para dar soporte a la visualización de la matriz '
        'de riesgo y registro detallado de excepciones y fallas intervinientes en el flujo operativo:'
    )
    
    doc.add_heading('A. Template: admin_risk_dashboard.html', level=3)
    doc.add_paragraph('Provee la representación gráfica del FMECA mediante el uso de Chart.js y componentes visuales. '
                      'Muestra los KPIs globales/locales (según el tenant u OID), la distribución del nivel de impacto '
                      '(bajo, moderado, severo), y un heatmap de RPN consolidado.')
    
    doc.add_heading('B. Template: admin_error_log.html', level=3)
    doc.add_paragraph('Pantalla que provee el catálogo interactivo de errores, utilizando componentes de filtrado paginado (Datatables y/o listados responsivos). '
                      'Permite interrogar los errores por modo de falla, severidad y rango de fechas. '
                      'Incorpora el pilar clave del sistema: la Vista de Perfil (Negocio vs. Técnico).')
                      
    doc.add_heading('C. Template: admin_error_detail.html', level=3)
    doc.add_paragraph('Expone el detalle unitario del error seleccionado, adaptándose según el parámetro de perfil recibido:'
                      '\n• Perfil Negocio: Oculta variables del entorno (CLOB/Traceback) y presenta un medidor de severidad con recomendaciones analíticas y funcionales de alto nivel.'
                      '\n• Perfil Técnico/Experto: Muestra la bitácora profunda del trace de Pila (Stacktrace) permitiendo al sysadmin diagnosticar exactamente en qué sentencia falló la operación, los parámetros crudos de ingreso, y la estructura subyacente de la excepción.')
    
    doc.add_heading('3. Inscripción al Maestro de Privilegios (Seed de Configuración)', level=2)
    doc.add_paragraph('Dado que estos templates brindan acceso a información corporativa que puede considerarse sensible (ej. fallas), '
                      'se procedió a registrar nuevos atributos al motor de roles y permisos dentro de las jerarquías de "AUDITORIA" y "SISTEMA":')
    
    doc.add_paragraph('view_risk_dashboard: Permite el acceso al Dashboard de Riesgos FMECA – Heatmap RPN, visualización de Modos de Falla y Trends de errores por módulo.', style=ul_style)
    doc.add_paragraph('view_error_log: Brinda el control para accesar la Consulta de Errores del Sistema (perfil negocio y experto).', style=ul_style)
    doc.add_paragraph('manage_mitigation_rules: Facilita el alta, baja y modificación de reglas FMECA para alertas y bloqueos del sistema.', style=ul_style)
    doc.add_paragraph('view_mitigation_history: Permite el seguimiento del historial de mitigaciones y acciones de respuesta automática tomadas frente a eventos críticos.', style=ul_style)
    
    doc.save(filepath)
    print("Documentos actualizados exitosamente.")

if __name__ == '__main__':
    append_to_word('c:\\Users\\marce\\Documents\\GitHub\\bibliotecaweb\\Informe_Tecnico_Detallado_FMECA_Resiliencia.docx')
